"""Generate Amgen Scholars proposal as a Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins (1 inch all around) ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

def set_style(paragraph, size_pt=11, bold=False, space_before=0, space_after=0):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after  = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size_pt)
        run.bold = bold

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_paragraph(
    "Quantifying Conformational Population Uncertainty in Cryo-EM Reconstructions "
    "of Human CFTR Using Gaussian Mixture Models and Deep Learning"
)
for run in title.runs:
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True
title.paragraph_format.space_after  = Pt(0)
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# ── Authors ────────────────────────────────────────────────────────────────────
authors = doc.add_paragraph("Minou Emmad, John F. Hunt")
set_style(authors, bold=False, space_after=6)

# ── Body paragraphs ────────────────────────────────────────────────────────────
paragraphs = [
    ("",
     "Cystic fibrosis (CF) is a life-threatening genetic disease affecting approximately "
     "35,000 people in the United States, caused by mutations in the Cystic Fibrosis "
     "Transmembrane Conductance Regulator (CFTR), a protein that acts as a chloride ion "
     "channel essential for lung function. Many CF mutations prevent CFTR from folding "
     "correctly, and drugs such as elexacaftor/tezacaftor/ivacaftor (Trikafta) have "
     "transformed treatment by correcting this defect (Veit et al., 2023). However, the "
     "structural mechanisms by which these drugs act remain incompletely understood, "
     "limiting development of next-generation therapies (Wang et al., 2022)."),

    ("",
     "Cryo-electron microscopy (cryo-EM) allows imaging of hundreds of thousands of "
     "individual protein molecules and computational sorting into distinct three-dimensional "
     "conformational classes. The relative sizes of these classes—the conformational "
     "populations—describe how a protein distributes across its possible shapes and how "
     "drugs shift that balance. Estimating these populations accurately is difficult because "
     "the sorting algorithms are imperfect and introduce systematic errors (Punjani & "
     "Fleet, 2023; Evans et al., 2025)."),

    ("",
     "This project builds on preliminary work comparing the standard approach—assigning "
     "each particle image to one class based on its highest score (a \"hard assignment\")—"
     "against a new probabilistic approach using Gaussian Mixture Models (GMMs). CryoSPARC, "
     "the software used to process CFTR datasets in the Hunt lab, computes a probability "
     "score for each particle belonging to each class. Rather than discarding this "
     "information, the GMM framework uses the full distribution of scores to estimate class "
     "overlap and mathematically correct population counts for misclassification. Applied "
     "to two CFTR datasets totaling 230,396 particles each, this produced stable corrected "
     "population estimates with bootstrap uncertainties of ±0.2%, and confirmed that the "
     "three conformational classes are distinguishable at the population level despite "
     "individual particle assignments being inherently noisy."),

    ("",
     "Going forward, I plan to implement methods that eliminate the need for discrete class "
     "assignments altogether. I will begin with cryoDRGN, a deep learning framework that "
     "uses a variational autoencoder (VAE) to map each particle image to a point in a "
     "continuous low-dimensional latent space representing its conformation (Zhong et al., "
     "2021). This allows the full conformational landscape to be explored without "
     "pre-specifying a number of classes. I will then evaluate related latent-space "
     "generative models for whether they provide more principled uncertainty estimates for "
     "interpreting CFTR conformational dynamics."),

    ("",
     "I have been working on this project since May 2026."),
]

for _, text in paragraphs:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)

# ── References ─────────────────────────────────────────────────────────────────
ref_header = doc.add_paragraph("References")
set_style(ref_header, bold=True, space_before=4, space_after=2)

refs = [
    "Elborn J.S. (2016) Cystic fibrosis. Lancet 388: 2519–2531.",
    "Punjani A. and Fleet D.J. (2023) 3D variability analysis: Resolving continuous "
    "flexibility and discrete heterogeneity from single particle cryo-EM. J. Struct. Biol. "
    "213: 107699.",
    "Veit G. et al. (2023) Elexacaftor/tezacaftor/ivacaftor corrects the folding defect "
    "of F508del-CFTR through a dual mechanism. Science Advances 9: eadg5580.",
    "Zhong E.D. et al. (2021) CryoDRGN: reconstruction of heterogeneous cryo-EM structures "
    "using neural networks. Nature Methods 18: 176–185.",
    "Evans L. et al. (2025) Counting particles could give wrong probabilities in "
    "Cryo-Electron Microscopy. bioRxiv 2025.03.27.644168.",
    "Wang C. et al. (2022) Mechanism of dual pharmacological correction and potentiation "
    "of human CFTR. bioRxiv 2022.10.10.510913.",
]

for ref in refs:
    r = doc.add_paragraph(ref)
    r.paragraph_format.space_before = Pt(0)
    r.paragraph_format.space_after  = Pt(3)
    r.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in r.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)

out = "amgen_proposal.docx"
doc.save(out)
print(f"Saved: {out}")

# Word count (body only)
body = " ".join(text for _, text in paragraphs)
words = len(body.split())
print(f"Body word count: {words}")
